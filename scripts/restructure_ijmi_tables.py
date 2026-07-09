"""
IJMI投稿原稿: 本文表を3つに制限し、超過分をSupplementary Tablesへ移動する。

本文に残す表: Table 1-3
補足表: S1 感度分析, S2 構造変化, S3 監査ドメイン
"""
from __future__ import annotations

import shutil
from copy import deepcopy
from datetime import date
from pathlib import Path

from docx import Document

PKG = Path(__file__).resolve().parents[1] / "submission_package_IJMI"
MANUSCRIPT = PKG / "IJMI_Paper1_Manuscript_IJMI_Submission.docx"
SUPP_FILE = PKG / "IJMI_Paper1_Supplementary_Tables.docx"
BACKUP_SUFFIX = date.today().strftime("%Y%m%d")


def backup(path: Path) -> Path:
    dest = path.with_name(f"{path.stem}_backup_{BACKUP_SUFFIX}{path.suffix}")
    shutil.copy2(path, dest)
    return dest


def find_paragraph(doc: Document, prefix: str):
    for para in doc.paragraphs:
        if para.text.strip().startswith(prefix):
            return para
    raise ValueError(f"Paragraph not found: {prefix!r}")


def table_after_caption(doc: Document, caption_para) -> int:
    """キャプション直後の tbl 要素に対応する doc.tables のインデックスを返す。"""
    body = doc.element.body
    children = list(body)
    cap_el = caption_para._element
    cap_idx = children.index(cap_el)
    for nxt in children[cap_idx + 1 :]:
        if nxt.tag.endswith("tbl"):
            tbl_el = nxt
            for i, table in enumerate(doc.tables):
                if table._tbl is tbl_el:
                    return i
            break
    raise ValueError("Table element not found after caption")


def remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def clone_table_after_paragraph(paragraph, source_table):
    """キャプション段落の直後に表を挿入する。"""
    new_tbl = deepcopy(source_table._tbl)
    paragraph._element.addnext(new_tbl)
    return new_tbl


def add_caption_and_table(doc: Document, caption_text: str, source_table) -> None:
    p = doc.add_paragraph()
    run = p.add_run(caption_text)
    run.bold = True
    clone_table_after_paragraph(p, source_table)
    doc.add_paragraph("")


def build_supplementary_tables(
    sens_doc: Document,
    structural_table,
    domains_table,
) -> None:
    """補足表DOCXを S1-S3 で再構築する。"""
    doc = Document()
    title = doc.add_paragraph()
    title.add_run("Supplementary Material").bold = True
    doc.add_paragraph("")

    s1_caption = find_paragraph(sens_doc, "Table S1.")
    s1_idx = table_after_caption(sens_doc, s1_caption)
    add_caption_and_table(doc, s1_caption.text.strip(), sens_doc.tables[s1_idx])

    add_caption_and_table(
        doc,
        "Table S2. Structural changes documented during extraction.",
        structural_table,
    )

    add_caption_and_table(
        doc,
        "Table S3. Pre-model audit domains for administrative healthcare data used in AI-oriented research.",
        domains_table,
    )

    doc.save(SUPP_FILE)


def patch_manuscript_text(doc: Document) -> None:
    replacements = {
        "Three sensitivity scenarios were specified.": (
            "Three sensitivity scenarios were specified (Supplementary Table S1)."
        ),
        "Three sensitivity scenarios were summarized.": (
            "Three sensitivity scenarios were summarized (Supplementary Table S1)."
        ),
        "Four additional structural changes were documented during full-panel extraction.": (
            "Four additional structural changes were documented during full-panel extraction "
            "(Supplementary Table S2)."
        ),
        "Instead, it identified measurable domains that can later be formalized:": (
            "Instead, it identified measurable domains that can later be formalized "
            "(Supplementary Table S3):"
        ),
    }
    for para in doc.paragraphs:
        text = para.text
        for old, new in replacements.items():
            if old in text and new not in text:
                para.text = text.replace(old, new)


def remove_main_table(doc: Document, caption_prefix: str) -> None:
    caption = find_paragraph(doc, caption_prefix)
    tbl_idx = table_after_caption(doc, caption)
    table = doc.tables[tbl_idx]
    remove_element(table._tbl)
    remove_element(caption._element)


def main() -> None:
    old_s1 = PKG / "IJMI_Paper1_Supplementary_Table_S1.docx"
    if not old_s1.exists():
        raise FileNotFoundError(old_s1)

    manuscript_backup = PKG / f"IJMI_Paper1_Manuscript_IJMI_Submission_backup_{BACKUP_SUFFIX}.docx"
    if not manuscript_backup.exists():
        backup(MANUSCRIPT)

    # 補足表のみ再構築（原稿は既に更新済みの場合はバックアップから表を取得）
    source_manuscript = manuscript_backup if manuscript_backup.exists() else MANUSCRIPT
    doc = Document(str(source_manuscript))
    sens_doc = Document(str(old_s1))

    cap4 = find_paragraph(doc, "Table 4.")
    cap5 = find_paragraph(doc, "Table 5.")
    tbl4 = doc.tables[table_after_caption(doc, cap4)]
    tbl5 = doc.tables[table_after_caption(doc, cap5)]

    build_supplementary_tables(sens_doc, tbl4, tbl5)

    # 原稿が未更新なら本文表も整理する
    if any(p.text.strip().startswith("Table 4.") for p in Document(str(MANUSCRIPT)).paragraphs):
        live = Document(str(MANUSCRIPT))
        remove_main_table(live, "Table 5.")
        remove_main_table(live, "Table 4.")
        patch_manuscript_text(live)
        live.save(MANUSCRIPT)

    print(f"Updated supplementary tables: {SUPP_FILE}")


if __name__ == "__main__":
    main()
