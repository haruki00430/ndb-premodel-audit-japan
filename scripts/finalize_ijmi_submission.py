"""
IJMI投稿前の最終修正: プレースホルダ削除・構造変化件数統一・日付更新。
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document

PKG = Path(__file__).resolve().parents[1] / "submission_package_IJMI"
MANUSCRIPT = PKG / "IJMI_Paper1_Manuscript_IJMI_Submission.docx"
COVER = PKG / "IJMI_Paper1_Cover_Letter.docx"
HIGHLIGHTS = PKG / "IJMI_Paper1_Highlights.docx"
SUMMARY = PKG / "IJMI_Paper1_Summary_Table.docx"

OLD_RESULTS_46 = (
    "Four additional structural changes were documented during full-panel extraction "
    "(Supplementary Table S2). LDL and HDL labels were shortened in releases No.6 to No.11, "
    "which required normalization and bidirectional matching to avoid missed extraction. "
    "HbA1C unit notation was removed from item labels in releases No.6 to No.11. "
    "In releases No.10 and No.11, file structures changed for mean-value data. "
    "Dental suppression rates were lower from FY2022 onward. "
    "These changes were recorded as data-structure findings rather than interpreted as model performance results."
)
NEW_RESULTS_46 = (
    "Nine undocumented structural changes were identified across releases No.1 to No.11. "
    "Supplementary Table S2 summarizes five extraction-relevant changes documented during "
    "full-panel extraction, including LDL and HDL label shortening in releases No.6 to No.11, "
    "which required normalization and bidirectional matching to avoid missed extraction; "
    "HbA1C unit notation removal from item labels in releases No.6 to No.11; "
    "mean-value file-structure changes in releases No.10 and No.11; "
    "the No.8 dental metric change; and lower dental suppression rates from FY2022 onward. "
    "These changes were recorded as data-structure findings rather than interpreted as model performance results."
)

HIGHLIGHT_OLD = (
    "Pre-model audit of NDB Open Data identified four undocumented structural changes"
)
HIGHLIGHT_NEW = (
    "Nine undocumented structural changes identified across all 11 NDB Open Data releases"
)

COVER_DATE_OLD = "8 July 2026"
COVER_DATE_NEW = "9 July 2026"


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def fix_manuscript() -> None:
    doc = Document(str(MANUSCRIPT))
    for para in list(doc.paragraphs):
        text = para.text.strip()
        if text == "Statement of significance.":
            remove_paragraph(para)
        elif text == OLD_RESULTS_46:
            para.text = NEW_RESULTS_46
    doc.save(MANUSCRIPT)


def replace_in_paragraphs(doc: Document, old: str, new: str) -> int:
    count = 0
    for para in doc.paragraphs:
        if old in para.text:
            para.text = para.text.replace(old, new)
            count += 1
        for cell in _iter_table_cells(doc):
            if old in cell.text:
                cell.text = cell.text.replace(old, new)
                count += 1
    return count


def _iter_table_cells(doc: Document):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield cell


def fix_highlights() -> None:
    doc = Document(str(HIGHLIGHTS))
    for para in doc.paragraphs:
        if HIGHLIGHT_OLD in para.text:
            para.text = para.text.replace(HIGHLIGHT_OLD, HIGHLIGHT_NEW)
    doc.save(HIGHLIGHTS)


def fix_cover_letter() -> None:
    doc = Document(str(COVER))
    for para in doc.paragraphs:
        if COVER_DATE_OLD in para.text:
            para.text = para.text.replace(COVER_DATE_OLD, COVER_DATE_NEW)
    doc.save(COVER)


def fix_summary_table() -> None:
    doc = Document(str(SUMMARY))
    # Summary table already says nine; ensure consistency if any 'four' remains
    for cell in _iter_table_cells(doc):
        if "four undocumented" in cell.text.lower():
            cell.text = cell.text.replace(
                "Four undocumented structural changes",
                "Nine undocumented structural changes",
            )
    doc.save(SUMMARY)


def main() -> None:
    fix_manuscript()
    fix_highlights()
    fix_cover_letter()
    fix_summary_table()
    print("DOCX fixes applied:", date.today())


if __name__ == "__main__":
    main()
